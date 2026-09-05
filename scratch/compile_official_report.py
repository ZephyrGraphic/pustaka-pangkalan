# -*- coding: utf-8 -*-
"""
Compilation Pipeline for Official Academic Report
Converts docs/laporan_utuh_pustaka_pangkalan.md into:
1. Standalone Academic LaTeX (.tex)
2. Polished Microsoft Word (.docx) with formal margins, typography, and styled tables
3. Native High-Resolution PDF (.pdf) via Word COM Engine
"""

import os
import sys
import subprocess
import win32com.client
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')

BASE_DIR = r"D:\CODEX-PROJECT\Perpustakaan Digital"
DOCS_DIR = os.path.join(BASE_DIR, "docs")
MD_PATH = os.path.join(DOCS_DIR, "laporan_utuh_pustaka_pangkalan.md")
TEX_PATH = os.path.join(DOCS_DIR, "laporan_utuh_pustaka_pangkalan.tex")
DOCX_PATH = os.path.join(DOCS_DIR, "laporan_utuh_pustaka_pangkalan.docx")
PDF_PATH = os.path.join(DOCS_DIR, "laporan_utuh_pustaka_pangkalan.pdf")

print("=== PIPELINE KOMPILASI LAPORAN SISTEM INFORMASI RESMI ===")

# 1. Generate standalone academic LaTeX source file (.tex)
print("\n[1/4] Mengompilasi naskah ke format LaTeX Standar Akademik (.tex)...")
tex_cmd = [
    "pandoc",
    MD_PATH,
    "-o", TEX_PATH,
    "--standalone",
    "--toc",
    "--toc-depth=3",
    "--number-sections",
    "-V", "documentclass=report",
    "-V", "geometry:a4paper,top=2.5cm,bottom=2.5cm,left=3cm,right=2.5cm",
    "-V", "fontsize=12pt",
    "-V", "linestretch=1.15"
]
subprocess.run(tex_cmd, check=True)
print(f"-> Berhasil menghasilkan file LaTeX: {TEX_PATH} ({os.path.getsize(TEX_PATH)} bytes)")

# 2. Generate Word DOCX via Pandoc
print("\n[2/4] Mengompilasi naskah ke format Microsoft Word (.docx)...")
docx_cmd = [
    "pandoc",
    MD_PATH,
    "-o", DOCX_PATH,
    "--toc",
    "--toc-depth=3",
    "--number-sections"
]
subprocess.run(docx_cmd, check=True)
print(f"-> Pandoc DOCX berhasil dibuat: {DOCX_PATH} ({os.path.getsize(DOCX_PATH)} bytes)")

# 3. Polish DOCX with Python-docx
print("\n[3/4] Melakukan penyempurnaan tipografi dan layout formal pada berkas DOCX...")
doc = Document(DOCX_PATH)

# Set formal academic margins: Top 2.5cm, Bottom 2.5cm, Left 3.0cm (binding margin), Right 2.5cm
for section in doc.sections:
    section.top_margin = Inches(0.984)     # ~2.5 cm
    section.bottom_margin = Inches(0.984)  # ~2.5 cm
    section.left_margin = Inches(1.181)    # ~3.0 cm (standar jilid resmi)
    section.right_margin = Inches(0.984)   # ~2.5 cm
    
    # Running Header
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Pustaka Pangkalan — Sistem Informasi Perpustakaan Digital Desa Pangkalan, Kec. Cikidang"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.name = "Times New Roman"
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.italic = True
        hp.runs[0].font.color.rgb = RGBColor(100, 100, 100)

# Format paragraph typography
for p in doc.paragraphs:
    # Justify body paragraphs
    if not p.text.startswith("#") and p.style.name.startswith("Normal"):
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            if not r.font.name:
                r.font.name = "Times New Roman"

# Format tables
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # Padding cell
            tcMar = OxmlElement('w:tcMar')
            for m in ['top', 'bottom']:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), '120')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            for m in ['left', 'right']:
                node = OxmlElement(f'w:{m}')
                node.set(qn('w:w'), '160')
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            
            # Header row styling (Forest Green #1E3A2F, white bold text)
            if i == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A2F"/>')
                tcPr.append(shading)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        r.font.name = "Times New Roman"
            else:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.font.name:
                            r.font.name = "Times New Roman"

doc.save(DOCX_PATH)
print(f"-> Penyempurnaan tipografi DOCX berhasil: {DOCX_PATH} ({os.path.getsize(DOCX_PATH)} bytes)")

# 4. Convert polished DOCX to PDF via Word COM
print("\n[4/4] Mengonversi DOCX ke Dokumen PDF Resmi beresolusi tinggi (Microsoft Word Engine)...")
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_com = word.Documents.Open(DOCX_PATH)
    doc_com.SaveAs(PDF_PATH, FileFormat=17) # 17 = wdFormatPDF
    doc_com.Close(False)
    word.Quit()
    print(f"-> Dokumen PDF resmi berhasil diekspor: {PDF_PATH} ({os.path.getsize(PDF_PATH)} bytes)")
except Exception as e:
    print(f"Catatan Word COM Automation: {e}")

print("\n=======================================================")
print("🎉 SELURUH LAPORAN SISTEM INFORMASI RESMI TELAH SELESAI!")
print(f"1. Naskah Markdown : {MD_PATH}")
print(f"2. Dokumen LaTeX   : {TEX_PATH}")
print(f"3. Dokumen Word    : {DOCX_PATH}")
print(f"4. Dokumen PDF     : {PDF_PATH}")
print("=======================================================")
